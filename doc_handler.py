from docx import Document
from copy import deepcopy
import re


class DocHandler:

	@staticmethod
	def insert(input_file, output_file, token, data_to_insert_list):
		"""

		:param input_file:
		:param output_file:
		:param token:
		:param data_to_insert_list:
		:return:
		"""
		data_to_insert_text = "\n".join(data_to_insert_list)

		try:

			document = Document(input_file)

			for paragraph in document.paragraphs:

				for run in paragraph.runs:

					if token in run.text:

						run.text = run.text.replace(token, data_to_insert_text)

			document.save(output_file)

			return output_file

		except Exception as ex:
			print("[ERROR]%s" % ex)
			return None

	@staticmethod
	def insert2(input_file, output_file, token, data_to_insert_list):
		"""

		:param input_file:
		:param output_file:
		:param token:
		:param data_to_insert_list:
		:return:
		"""
		try:

			document = Document(input_file)

			for paragraph in document.paragraphs:

				if token in paragraph.text:

					old_paragraph = deepcopy(paragraph)

					paragraph.clear()

					for data_to_insert in data_to_insert_list:
						new_paragraph = document.add_paragraph()
						new_paragraph_run = new_paragraph.add_run()
						new_paragraph_run.add_text(data_to_insert)

			document.save(output_file)

			return output_file

		except Exception as ex:
			print("[ERROR]%s" % ex)
			return None
